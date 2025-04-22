import docx
from tkinter import Tk, Label, Entry, Checkbutton, BooleanVar, Button, ttk, messagebox
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from pathlib import Path
from docx.oxml.ns import qn
from docx.shared import Inches

def gerar_nome_unico(diretorio, nome_base, extensao):
    """
    Gera um nome de arquivo único adicionando um número incremental
    """
    contador = 0
    nome_arquivo = f"{nome_base}{extensao}"
    caminho_completo = os.path.join(diretorio, nome_arquivo)
    
    while os.path.exists(caminho_completo):
        contador += 1
        nome_arquivo = f"{nome_base}_{contador}{extensao}"
        caminho_completo = os.path.join(diretorio, nome_arquivo)
    
    return caminho_completo

def adicionar_rodape_com_imagem(doc, caminho_imagem, largura=Inches(6.1)):
    """Adiciona uma imagem centralizada no rodapé do documento"""
    try:
        section = doc.sections[0]
        footer = section.footer
        
        # Limpa o rodapé existente de forma segura
        for p in footer.paragraphs:
            p.clear()
        
        # Cria um novo parágrafo centralizado
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Adiciona a imagem
        run = paragraph.add_run()
        run.add_picture(caminho_imagem, width=largura)
        
        # Adiciona espaço após a imagem (opcional)
        paragraph.space_after = Pt(10)
        
    except Exception as e:
        messagebox.showerror("Erro", f"Não foi possível adicionar rodapé:\n{str(e)}")

# Dados fixos

###SUBSTITUIR DADOS PELOS REAIS
EMPRESA = "EMPRESA SHOW DE BOLA"
CNPJ_EMPRESA = "01.234.567/0001-89"
CIDADE = "Cidade Legal"
VALOR_SERVICO = 150.00
VALOR_PASSAGEM = 20.00
DESCONTO_DAS = 80.90

# Dados dos funcionários

###SUBSTITUIR DADOS PELOS REAIS
funcionarios = {
    'FUNCIONÁRIO 1': {
        'cnpj': '12.345.678/0001-90',
        'servico': 'serviço geral'
    },
    'FUNCIONARIO 2': {
        'cnpj': '98.765.432/0001-10',
        'servico': 'serviço específico'
    }
}

def formatar_moeda_br(valor):
    """Formata valores float no padrão R$ brasileiro (1.234,56)"""
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def gerar_recibo():
    funcionario_selecionado = combo_funcionario.get()
    aplicar_desconto = var_desconto.get()
    
    if funcionario_selecionado == "Outro":
        nome_funcionario = entry_nome.get()
        cnpj_funcionario = entry_cnpj.get()
        servico = entry_servico.get()
        
        if not all([nome_funcionario, cnpj_funcionario, servico]):
            messagebox.showerror("Erro", "Preencha todos os campos para 'Outro'!")
            return
            
        func = {
            'nome': nome_funcionario,
            'cnpj': cnpj_funcionario,
            'servico': servico
        }
    elif funcionario_selecionado in funcionarios:
        func = funcionarios[funcionario_selecionado]
        func['nome'] = funcionario_selecionado
    else:
        messagebox.showerror("Erro", "Selecione uma opção válida!")
        return

    # Cálculos
    total = VALOR_SERVICO + VALOR_PASSAGEM
    if aplicar_desconto:
        total -= DESCONTO_DAS

    # Criar documento
    doc = Document()

    adicionar_rodape_com_imagem(doc, "rodape_img.png", largura=Inches(6.1))
    
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    
    # Espaçamento inicial
    doc.add_paragraph()
    
    # Título RECIBO
    titulo = doc.add_paragraph("RECIBO")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    titulo.runs[0].font.size = Pt(18)
    titulo.runs[0].bold = True
    
    # Espaçamento
    for _ in range(2):
        doc.add_paragraph()
    
    # Texto principal
    texto_principal = doc.add_paragraph()
    texto_principal.add_run(f"Recebi de {EMPRESA} sob o CNPJ nº {CNPJ_EMPRESA} a quantia de {formatar_moeda_br(total)} referente a {func['servico']}.")
    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    texto_principal.runs[0].font.size = Pt(12)
    
    # Espaçamento
    doc.add_paragraph()
    
    # Detalhes dos valores
    detalhes = doc.add_paragraph()
    detalhes.add_run(f"• Serviço: {formatar_moeda_br(VALOR_SERVICO)}\n")
    detalhes.add_run(f"• Passagem: {formatar_moeda_br(VALOR_PASSAGEM)}\n")
    detalhes.runs[0].bold = True
    detalhes.runs[1].bold = True
    
    if aplicar_desconto:
        detalhes.add_run(f"• Desconto DAS MEI: - {formatar_moeda_br(DESCONTO_DAS)}\n")
        detalhes.runs[2].bold = True
    
    # Total
    doc.add_paragraph()
    total_para = doc.add_paragraph(f"Total: {formatar_moeda_br(total)}")
    total_para.runs[0].bold = True
    
    # Espaçamento
    for _ in range(3):
        doc.add_paragraph()
    
    # Data e local
    meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
             'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    hoje = datetime.now()
    data_por_extenso = f"{CIDADE}, {hoje.day} de {meses[hoje.month-1]} de {hoje.year}."
    data = doc.add_paragraph(data_por_extenso)
    data.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Espaçamento
    for _ in range(4):
        doc.add_paragraph()
    
    # Assinatura
    linha_assinatura = doc.add_paragraph("______________________________________________________")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Nome e CNPJ
    assinatura = doc.add_paragraph(f"{func['nome']}\nCNPJ {func['cnpj']}")
    assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Configurar diretório de saída

    ###SUBSTITUIR DADOS PELOS REAIS
    OUTPUT_FOLDER = r"caminho no seu computador"
    
    try:
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    except Exception as e:
        messagebox.showerror("Erro", f"Não foi possível acessar o diretório:\n{str(e)}")
        return

    # Gerar nome do arquivo mantendo acentos
    nome_base = func['nome'].split()[0] if funcionario_selecionado != "Outro" else "Outro"
    nome_arquivo_base = f"Recibo_{nome_base}_{datetime.now().strftime('%Y%m%d')}"

    # Gerar caminho único
    caminho_completo = gerar_nome_unico(OUTPUT_FOLDER, nome_arquivo_base, ".docx")

    try:
        doc.save(caminho_completo)
        messagebox.showinfo("Sucesso", f"Recibo gerado:\n{caminho_completo}")
    except Exception as e:
        messagebox.showerror("Erro", f"Falha ao salvar:\n{str(e)}")

def mostrar_campos_outro(event):
    if combo_funcionario.get() == "Outro":
        frame_outro.grid(row=3, column=0, columnspan=2, padx=10, pady=5, sticky="we")
    else:
        frame_outro.grid_remove()

# Interface
root = Tk()
root.title("Gerador de Recibos")
root.geometry("500x400")

var_desconto = BooleanVar(value=False)

# Widgets principais
Label(root, text="Funcionário:").grid(row=0, column=0, padx=10, pady=5, sticky="e")
combo_funcionario = ttk.Combobox(root, values=list(funcionarios.keys()) + ["Outro"], state="readonly")
combo_funcionario.grid(row=0, column=1, padx=10, pady=5, sticky="we")
combo_funcionario.set(list(funcionarios.keys())[0])
combo_funcionario.bind("<<ComboboxSelected>>", mostrar_campos_outro)

# Checkbox para desconto
Checkbutton(root, text="Aplicar desconto DAS MEI (mensal)", variable=var_desconto).grid(row=1, column=0, columnspan=2, pady=5)

# Frame para campos "Outro"
frame_outro = ttk.LabelFrame(root, text="Dados para 'Outro'")
frame_outro.grid(row=3, column=0, columnspan=2, padx=10, pady=5, sticky="we")
frame_outro.grid_remove()  # Inicia oculto

Label(frame_outro, text="Nome Completo:").grid(row=0, column=0, padx=5, pady=2, sticky="e")
entry_nome = Entry(frame_outro)
entry_nome.grid(row=0, column=1, padx=5, pady=2, sticky="we")

Label(frame_outro, text="CNPJ/CPF:").grid(row=1, column=0, padx=5, pady=2, sticky="e")
entry_cnpj = Entry(frame_outro)
entry_cnpj.grid(row=1, column=1, padx=5, pady=2, sticky="we")

Label(frame_outro, text="Serviço Prestado:").grid(row=2, column=0, padx=5, pady=2, sticky="e")
entry_servico = Entry(frame_outro)
entry_servico.grid(row=2, column=1, padx=5, pady=2, sticky="we")

# Botão gerar
Button(root, text="Gerar Recibo", command=gerar_recibo).grid(row=4, column=0, columnspan=2, pady=20)

# Ajustes de layout
root.columnconfigure(1, weight=1)
frame_outro.columnconfigure(1, weight=1)

root.mainloop()
